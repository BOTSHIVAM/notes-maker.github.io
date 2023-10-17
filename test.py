import streamlit as st
import pyaudio
import wave
import speech_recognition as sr
import spacy
import docx
import os
import base64

# Create a directory to store history files
if not os.path.exists("history"):
    os.makedirs("history")

def main():
    st.title("Audio to Summary App")

    CHUNK = 1024
    FORMAT = pyaudio.paInt16
    CHANNELS = 1
    RATE = 44100
    OUTPUT_FILENAME = "output.wav"

    recording = False  # Flag to indicate whether recording is in progress

    if not recording:
        if st.button("Start Recording"):
            recording = True
            st.write("Recording...")

            audio = pyaudio.PyAudio()
            stream = audio.open(format=FORMAT, channels=CHANNELS, rate=RATE, input=True, frames_per_buffer=CHUNK)

            frames = []

            while recording:
                data = stream.read(CHUNK)
                frames.append(data)

            st.write("Recording finished.")
            st.write("Recognizing...")

            stream.stop_stream()
            stream.close()
            audio.terminate()

            with wave.open(OUTPUT_FILENAME, 'wb') as wf:
                wf.setnchannels(CHANNELS)
                wf.setsampwidth(audio.get_sample_size(FORMAT))
                wf.setframerate(RATE)
                wf.writeframes(b''.join(frames))

            recognizer = sr.Recognizer()

            audio_file = "output.wav"
            with sr.AudioFile(audio_file) as source:
                audio_data = recognizer.record(source)
                text = recognizer.recognize_google(audio_data)

            st.write("Recognized text:", text)

            nlp = spacy.load("en_core_web_sm")
            doc = nlp(text)

            # Autocorrect the recognized text
            corrected_text = " ".join([token.text for token in doc])

            def create_bullet_points(text):
                sentences = text.split('.')
                bullet_points = [f"• {sentence.strip()}" for sentence in sentences if sentence.strip()]
                return '\n'.join(bullet_points)

            summary = create_bullet_points(corrected_text)

            st.write("Summary:")
            st.write(summary)

            doc = docx.Document()
            doc.add_paragraph("Summary:")
            doc.add_paragraph(summary)

            doc_file = os.path.join("history", f"summary_{len(os.listdir('history')) + 1}.docx")
            doc.save(doc_file)

            recording = False  # Reset the recording flag

            # Provide a link to download the summary document
            st.markdown(get_binary_file_downloader_html(doc_file, 'Download Summary Document'))

    # Display the history section in the sidebar
    st.sidebar.header("History")
    history_files = [file for file in os.listdir("history") if file.endswith(".docx")]
    selected_file = st.sidebar.selectbox("Select a file from history", history_files, key="history_select", index=0)

    if selected_file:
        doc_file = os.path.join("history", selected_file)
        st.sidebar.markdown(get_binary_file_downloader_html(doc_file, 'Download'))

def get_binary_file_downloader_html(bin_file, label):
    with open(bin_file, 'rb') as f:
        data = f.read()
    b64 = base64.b64encode(data).decode()
    href = f'<a href="data:application/octet-stream;base64,{b64}" download="{os.path.basename(bin_file)}">{label}</a>'
    return href

if __name__ == "__main__":
    main()
